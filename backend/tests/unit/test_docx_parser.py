from io import BytesIO

import pytest
from docx import Document

from app.models.import_models import ImportRequest
from app.services.parser_service import DocxParser, EmptyContentError, ParserError


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def docx_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_request(source_bytes: bytes, filename: str = "document.docx") -> ImportRequest:
    return ImportRequest(filename=filename, mime_type=DOCX_MIME, source_bytes=source_bytes)


def test_docx_parser_extracts_simple_paragraphs() -> None:
    document = Document()
    document.add_paragraph("Erster Absatz")
    document.add_paragraph("Zweiter Absatz")
    source = docx_bytes(document)

    parsed = DocxParser().parse(make_request(source))

    assert parsed.text == "Erster Absatz\n\nZweiter Absatz\n"
    assert parsed.parser_name == "docx-parser"
    assert parsed.metadata["parser_name"] == "docx-parser"
    assert parsed.metadata["mime_type"] == DOCX_MIME
    assert parsed.metadata["source_filename"] == "document.docx"
    assert parsed.metadata["byte_size"] == len(source)
    assert parsed.metadata["paragraph_count"] == 2
    assert parsed.metadata["table_count"] == 0


def test_docx_parser_converts_table_to_markdown_table() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Wert"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"

    parsed = DocxParser().parse(make_request(docx_bytes(document)))

    assert "| Name | Wert |" in parsed.text
    assert "| --- | --- |" in parsed.text
    assert "| Alpha | 42 |" in parsed.text
    assert parsed.metadata["table_count"] == 1


def test_docx_parser_converts_headings_to_markdown_headings() -> None:
    document = Document()
    document.add_heading("Haupttitel", level=1)
    document.add_heading("Untertitel", level=2)
    document.add_paragraph("Inhalt")

    parsed = DocxParser().parse(make_request(docx_bytes(document)))

    assert parsed.text == "# Haupttitel\n\n## Untertitel\n\nInhalt\n"
    assert parsed.metadata["paragraph_count"] == 3


def test_docx_parser_rejects_empty_docx() -> None:
    document = Document()

    with pytest.raises(EmptyContentError, match="does not contain importable content"):
        DocxParser().parse(make_request(docx_bytes(document), filename="empty.docx"))


def test_docx_parser_rejects_corrupted_docx() -> None:
    with pytest.raises(ParserError, match="could not be opened"):
        DocxParser().parse(make_request(b"not a valid docx", filename="broken.docx"))
